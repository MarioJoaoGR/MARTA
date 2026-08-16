"""Converte apresentacao/diagramas.html num .docx editável.

Porquê: o HTML é bom para ver e mau para editar. O Word é o contrário. Este
script mantém o HTML como fonte única e gera o .docx a partir dele, para não
haver duas versões do texto a divergir.

Os diagramas em SVG são convertidos a PNG com o `qlmanage` do macOS, que é o
único conversor disponível sem instalar nada. Ele só produz miniaturas
quadradas e estica o desenho para as encher, por isso o SVG é primeiro
recentrado numa tela quadrada; a imagem final é recortada pelos pixéis com
tinta.

    python scripts/html_para_docx.py
    python scripts/html_para_docx.py --html X.html --out Y.docx
"""
from __future__ import annotations

import argparse
import html
import os
import re
import subprocess
import sys
import tempfile

MONO = "Consolas"


def svg_para_png(svg: str, css: str, dest: str, tmp: str, idx: int) -> str | None:
    """SVG isolado -> PNG recortado. Devolve o caminho, ou None se falhar."""
    m = re.search(r'viewBox="0 0 (\d+) (\d+)"', svg)
    if not m:
        return None
    w, h = int(m.group(1)), int(m.group(2))
    lado = max(w, h)
    mx, my = -(lado - w) // 2, -(lado - h) // 2
    # O CSS vive na página, não no SVG: sem o embutir, o desenho sai sem cores.
    corpo = svg.replace(
        "<svg ", f'<svg xmlns="http://www.w3.org/2000/svg" width="{lado}" height="{lado}" ', 1)
    corpo = re.sub(r'viewBox="0 0 \d+ \d+"', f'viewBox="{mx} {my} {lado} {lado}"', corpo, count=1)
    corpo = corpo.replace(
        ">", f'><style>{css}</style>'
             f'<rect x="{mx}" y="{my}" width="{lado}" height="{lado}" fill="#ffffff"/>', 1)
    src = os.path.join(tmp, f"f{idx}.svg")
    with open(src, "w", encoding="utf-8") as f:
        f.write('<?xml version="1.0" encoding="UTF-8"?>\n' + corpo)
    subprocess.run(["qlmanage", "-t", "-s", "2200", "-o", tmp, src],
                   capture_output=True, timeout=120)
    bruto = src + ".png"
    if not os.path.exists(bruto):
        return None
    from PIL import Image, ImageChops
    im = Image.open(bruto).convert("RGB")
    dif = ImageChops.difference(im, Image.new("RGB", im.size, "white")).convert("L")
    bb = dif.point(lambda v: 255 if v > 8 else 0).getbbox()
    if not bb:
        return None
    pad = 20
    im.crop((max(0, bb[0] - pad), max(0, bb[1] - pad),
             min(im.size[0], bb[2] + pad), min(im.size[1], bb[3] + pad))).save(dest)
    return dest


def limpo(s: str, manter_bordas: bool = False) -> str:
    """Texto de um fragmento HTML, sem tags e com o espaço normalizado.

    ``manter_bordas`` preserva um espaço no início/fim quando o original o
    tinha: sem isso, o espaço entre ``</strong>`` e a palavra seguinte
    desaparecia e o texto saía colado ("O que faz:lê o texto").
    """
    s = re.sub(r"<br\s*/?>", " ", s)
    s = re.sub(r"<[^>]+>", "", s)
    s = re.sub(r"\s+", " ", html.unescape(s))
    return s if manter_bordas else s.strip()


def corridos(par, frag: str) -> None:
    """Escreve o fragmento no parágrafo preservando negrito, itálico e código."""
    frag = re.sub(r"<br\s*/?>", " ", frag)
    pecas: list[tuple[str, str]] = []
    for pedaco in re.split(r"(<(?:strong|b|em|i|code)>.*?</(?:strong|b|em|i|code)>)", frag, flags=re.S):
        if not pedaco:
            continue
        m = re.match(r"<(strong|b|em|i|code)>(.*?)</\1>", pedaco, re.S)
        if m:
            pecas.append((m.group(1), limpo(m.group(2))))
        else:
            pecas.append(("", limpo(pedaco, manter_bordas=True)))
    # o parágrafo não começa nem acaba com espaço; no meio, mantém-se
    while pecas and not pecas[0][1].strip():
        pecas.pop(0)
    while pecas and not pecas[-1][1].strip():
        pecas.pop()
    if not pecas:
        return
    pecas[0] = (pecas[0][0], pecas[0][1].lstrip())
    pecas[-1] = (pecas[-1][0], pecas[-1][1].rstrip())
    for tag, txt in pecas:
        if not txt:
            continue
        r = par.add_run(txt)
        if tag in ("strong", "b"):
            r.bold = True
        elif tag in ("em", "i"):
            r.italic = True
        elif tag == "code":
            r.font.name = MONO
            r.font.size = Pt(9.5)


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--html", default="apresentacao/diagramas.html")
    ap.add_argument("--out", default="apresentacao/diagramas.docx")
    args = ap.parse_args()

    doc_html = open(args.html, encoding="utf-8").read()
    css = re.search(r"<style>(.*?)</style>", doc_html, re.S).group(1)
    corpo = doc_html[doc_html.index('<div class="wrap">'):]

    d = Document()
    sec = d.sections[0]
    largura_util = sec.page_width - sec.left_margin - sec.right_margin
    normal = d.styles["Normal"].font
    normal.name, normal.size = "Calibri", Pt(11)

    tmp = tempfile.mkdtemp(prefix="h2docx_")
    n_svg = 0
    # Percorre os blocos de topo pela ordem em que aparecem.
    blocos = re.finditer(
        r"<h1[^>]*>(.*?)</h1>|<h2[^>]*>(.*?)</h2>|<h3[^>]*>(.*?)</h3>|<h4[^>]*>(.*?)</h4>"
        r"|<p class=\"sub\">(.*?)</p>|<p>(.*?)</p>|<pre><code>(.*?)</code></pre>"
        r"|<table>(.*?)</table>|<ul>(.*?)</ul>|<ol>(.*?)</ol>"
        r"|(<svg .*?</svg>)|<figcaption>(.*?)</figcaption>"
        r"|<div class=\"flag\">(.*?)</div>", corpo, re.S)

    for m in blocos:
        h1, h2, h3, h4, sub, p, pre, tab, ul, ol, svg, cap, flag = m.groups()
        if h1:
            d.add_heading(limpo(h1), level=1)
        elif h2:
            d.add_heading(re.sub(r"^\d+\s*", "", limpo(h2)), level=2)
        elif h3 or h4:
            d.add_heading(limpo(h3 or h4), level=3)
        elif sub:
            par = d.add_paragraph()
            r = par.add_run(limpo(sub))
            r.italic = True
            r.font.color.rgb = RGBColor(0x5F, 0x5F, 0x5F)
        elif p:
            corridos(d.add_paragraph(), p)
        elif pre:
            txt = html.unescape(re.sub(r"<[^>]+>", "", pre))
            for linha in txt.strip("\n").split("\n"):
                par = d.add_paragraph()
                par.paragraph_format.space_after = Pt(0)
                r = par.add_run(linha)
                r.font.name, r.font.size = MONO, Pt(9)
        elif tab:
            linhas = re.findall(r"<tr>(.*?)</tr>", tab, re.S)
            if not linhas:
                continue
            cels = [re.findall(r"<t[hd][^>]*>(.*?)</t[hd]>", ln, re.S) for ln in linhas]
            ncol = max(len(c) for c in cels)
            t = d.add_table(rows=0, cols=ncol)
            t.style = "Table Grid"
            for i, linha in enumerate(cels):
                celulas = t.add_row().cells
                for j in range(ncol):
                    celulas[j].text = ""
                    if j < len(linha):
                        par = celulas[j].paragraphs[0]
                        corridos(par, linha[j])
                        if i == 0 or "<th" in linhas[i]:
                            for r in par.runs:
                                r.bold = True
        elif ul or ol:
            estilo = "List Bullet" if ul else "List Number"
            for li in re.findall(r"<li>(.*?)</li>", ul or ol, re.S):
                li = re.sub(r"<ul>.*?</ul>", "", li, flags=re.S)
                corridos(d.add_paragraph(style=estilo), li)
        elif svg:
            n_svg += 1
            png = os.path.join(tmp, f"fig{n_svg}.png")
            if svg_para_png(svg, css, png, tmp, n_svg):
                d.add_picture(png, width=largura_util)
                d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                d.add_paragraph(f"[diagrama {n_svg}: falhou a conversão]")
        elif cap:
            par = d.add_paragraph()
            r = par.add_run(limpo(cap))
            r.italic = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x5F, 0x5F, 0x5F)
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif flag:
            for sub_p in re.findall(r"<p>(.*?)</p>", flag, re.S):
                par = d.add_paragraph()
                par.paragraph_format.left_indent = Inches(0.3)
                corridos(par, sub_p)
                for r in par.runs:
                    if not r.bold:
                        r.font.color.rgb = RGBColor(0x8B, 0x2B, 0x24)

    os.makedirs(os.path.dirname(args.out) or ".", exist_ok=True)
    d.save(args.out)
    print(f"{args.out}  ({n_svg} diagramas, {len(d.paragraphs)} parágrafos)")
    return 0


if __name__ == "__main__":
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
    sys.exit(main())
